#!/usr/bin/env python3
"""Export docs/USER-GUIDE.md as a Word document you can finish by hand.

    python3 tools/make_guide_docx.py     ->  docs/SmartNoteAI-UserGuide.docx

Everything lands as real Word objects: Heading 1/2/3, List Bullet, native
tables, inline pictures (the cropped assets/guide/*.jpg), Intense Quote for
callouts. Page size A4, Helvetica-ish (Arial) at 10.5/1.4.
"""
import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD = os.path.join(ROOT, 'docs/USER-GUIDE.md')
IMG = os.path.join(ROOT, 'assets/guide')
KOFI = os.path.join(ROOT, 'assets/kofi-qr.png')
DEVICE = os.path.join(IMG, 'cover-device.png')
OUT = os.path.join(ROOT, 'docs/SmartNoteAI-UserGuide.docx')

GREY = RGBColor(0x59, 0x59, 0x59)


def demark(t):
    t = re.sub(r'\[IMAGE: [^\]]+\]', '', t)
    t = t.replace('`', '')
    t = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', t)
    return t.strip()


def runs(par, text):
    """Write text into a paragraph, honouring **bold** and *italic*."""
    for chunk in re.split(r'(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*))', text):
        if not chunk:
            continue
        if chunk.startswith('**') and chunk.endswith('**'):
            par.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith('*') and chunk.endswith('*'):
            par.add_run(chunk[1:-1]).italic = True
        else:
            par.add_run(chunk)


def style_document(doc):
    st = doc.styles['Normal']
    st.font.name = 'Arial'
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.25
    for name, size, before in (('Heading 1', 20, 26), ('Heading 2', 14, 18),
                               ('Heading 3', 11.5, 14)):
        s = doc.styles[name]
        s.font.name = 'Arial'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(6)
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    for attr, v in (('left_margin', 2.6), ('right_margin', 2.6),
                    ('top_margin', 2.4), ('bottom_margin', 2.2)):
        setattr(sec, attr, Cm(v))


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SUPERNOTE PLUGIN')
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('SmartNote AI')
    r.font.size = Pt(34)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('User Guide')
    r.font.size = Pt(18)
    if os.path.exists(DEVICE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(DEVICE, height=Cm(11))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Your Supernote in the age of AI; privacy first.')
    r.font.size = Pt(12)
    for label, value in (('AUTHOR', 'AgP42'),
                         ('SOURCES AND ISSUES', 'github.com/AgP42/SN-Plugin-SmartNoteAI'),
                         ('SUPPORT THE PROJECT', 'ko-fi.com/agp42')):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + '   ')
        r.font.size = Pt(8)
        r.bold = True
        r.font.color.rgb = GREY
        r = p.add_run(value)
        r.font.size = Pt(10.5)
    if os.path.exists(KOFI):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(KOFI, width=Cm(3))
    doc.add_page_break()


def add_table(doc, rows):
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Light Grid Accent 1'
    for i, (left, right) in enumerate(rows):
        cells = t.add_row().cells
        for cell, text, bold in ((cells[0], left, True), (cells[1], right, False)):
            cell.text = ''
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            runs(par, text)
            for r in par.runs:
                r.font.size = Pt(9.5)
                if bold or i == 0:
                    r.bold = True
    doc.add_paragraph()


def build():
    doc = Document()
    style_document(doc)
    cover(doc)

    lines = open(MD, encoding='utf-8').read().split('\n')
    start = next(i for i, l in enumerate(lines) if l.startswith('# 1.'))
    table, para, in_code, code = [], [], False, []

    def flush_par():
        if para:
            runs(doc.add_paragraph(), ' '.join(para))
            para.clear()

    def flush_table():
        if table:
            add_table(doc, table[:])
            table.clear()

    for raw in lines[start:]:
        l = raw.rstrip()
        if l.startswith('```'):
            flush_par(); flush_table()
            in_code = not in_code
            if not in_code and code:
                p = doc.add_paragraph()
                r = p.add_run('\n'.join(code))
                r.font.name = 'Consolas'
                r.font.size = Pt(8)
                code.clear()
            continue
        if in_code:
            code.append(l)
            continue
        if '[IMAGE:' in l:
            flush_par(); flush_table()
            for name in re.findall(r'\[IMAGE: ([^\]]+)\]', l):
                path = os.path.join(IMG, name)
                if not os.path.exists(path):
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Cm(12.5))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(name)
                r.font.size = Pt(8)
                r.font.color.rgb = GREY
            continue
        if l.startswith('# ') or l.startswith('## ') or l.startswith('### '):
            flush_par(); flush_table()
            level = 1 if l.startswith('# ') else (2 if l.startswith('## ') else 3)
            title = demark(l.lstrip('#').strip())
            if level == 1:
                doc.add_page_break()
            doc.add_heading(title, level=level)
            continue
        if l.startswith('|'):
            flush_par()
            cells = [demark(c.strip()) for c in l.strip('|').split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            if len(cells) >= 2:
                table.append((cells[0], ' '.join(cells[1:]).strip()))
            continue
        if l.startswith('- ') or l.startswith('* '):
            flush_par(); flush_table()
            runs(doc.add_paragraph(style='List Bullet'), demark(l[2:].strip()))
            continue
        if re.match(r'^\d+\. ', l):
            flush_par(); flush_table()
            runs(doc.add_paragraph(style='List Number'), demark(l.split('. ', 1)[1]))
            continue
        if l.startswith('> '):
            flush_par(); flush_table()
            p = doc.add_paragraph(style='Intense Quote')
            runs(p, demark(l[2:].strip()))
            continue
        if l.strip() in ('', '---'):
            flush_par(); flush_table()
            continue
        para.append(demark(l.strip()))
    flush_par(); flush_table()
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
